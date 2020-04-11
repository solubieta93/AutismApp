from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from Database.manager import Manager
import matplotlib.dates as dt
import datetime
from datetime import timedelta
import math
import xlsxwriter
import os
import re

# import calendar
# from datetime import timedelta
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn

# TODO: probar q funcione bien el creado de carpetas con los nuevos nombres


def sorted_aphanumeric(data):
    convert = lambda text: int(text) if text.isdigit() else text.lower()
    alphanum_key = lambda key: [convert(c) for c in re.split('([0-9]+)', key)]
    return sorted(data, key=alphanum_key)


def list_sto(m, table_name, medicaid_id, name):
    sto_list = m.sto_data_string(table_name, str(medicaid_id), str(name))[2]
    print 'sto list ' + str(sto_list)
    mastered = False
    current_index = 1
    sto_string = ''
    status = 'In progress'

    for i, sto in enumerate(sto_list):
        if sto[0]:
            s = str(sto[3]).split('(M')
            sto_string = sto_string + 'Sto' + str(i + 1) + '(M' + str(s[1]) + ' '
            current_index = i + 1
            mastered = True
    if current_index < len(sto_list):
        if mastered:
            date_initial = sto_list[current_index-1][1]
            month = dt.num2date(dt.datestr2num(date_initial)).month
            year = dt.num2date(dt.datestr2num(date_initial)).year
            if month + 1 > 12:
                year += 1
                date_initial = '1/1/' + str(year)
            else:
                date_initial = '{}/1/{}'.format(str(month + 1), str(year))
            sto_string = sto_string + 'Currently: STO ' + str(current_index + 1) + '(Initiated ' + str(
                date_initial) + ') '
        else:
            sto_string = sto_string + 'Currently: STO ' + str(current_index) + ' '
    else:
        # aqui preguntar por el LTO
        print str(name)
        lto = m.lto_data(table_name, str(medicaid_id), str(name))
        print 'lto'
        print lto
        if lto[0]:
            if lto[2][3]:
                status = 'MET'
            else:
                status = 'LTO in progress'
    return sto_string, status


def average_met(table_name, medicaid_id, name_prog_beh):
    m = Manager()
    print 'estoy en average'

    b, data_kid = m.list_dates(table_name, medicaid_id, name_prog_beh)
    average = 0
    if b:
        print 'entre a b'
        list_date = data_kid.keys()
        print 'AVERAGE'
        print 'list date'
        print list_date

        list_data = data_kid.values()
        count = -4
        if len(list_date)< 4:
            count = 0- len(list_date)
        list_valid_data = []
        valid_month = dt.num2date(dt.datestr2num(list_date[-1])).month
        while count is not 0:
            print 'count ' + str(count)
            m = dt.num2date(dt.datestr2num(list_date[count])).month

            if m == valid_month and list_data[count][0] is not -1:
                list_valid_data.append(list_data[count][0])
            count += 1
        if len(list_valid_data) > 0:
            average = math.fsum(list_valid_data) / len(list_valid_data)
    return average


def create_monthly(username, medicaid_id=None, kid_name=None):
    m = Manager()
    # username = 'yessica'
    # date = datetime.datetime.now() '10/31/2018'
    # puedo saber el ultimo dia del mes q quiera con calendar.monthrange(year, mes)[1]
    # today = datetime.datetime.now()
    # calendar.monthrange(today.year, today.month)[1])

    ultimo_dia_mes_anterior = datetime.date.today().replace(day=1) + datetime.timedelta(days=-1)
    date = "%s/%s/%s" % (ultimo_dia_mes_anterior.month, ultimo_dia_mes_anterior.day, ultimo_dia_mes_anterior.year)
    my_date = datetime.datetime.now()
    my_date = my_date.replace(day=1) - timedelta(days=1)
    my_date = my_date.strftime("%B") + ' ' + str(my_date.year)

    if medicaid_id is not None:
        list_kids = [str(medicaid_id) + '-->' + str(kid_name)]
        list_medicaid = [medicaid_id]
    else:
        list_kids = m.list_kids(username)
        list_medicaid = m.list_medicaid_kid(username)

    for j, medicaid_id in enumerate(list_medicaid):
        document = Document()

        name_kid = str(list_kids[j].split('-->')[1])

        document.add_heading('Measurable and Progress from maladaptive behaviors.', level=1)
        records, mastered, baseline_date = create_data_table(m, 'behaviors', medicaid_id)
        fill_table(document, 'behaviors', records, baseline_date, date)
        fill_with_graphics(document, 'behaviors', mastered, my_date, name_kid)

        document.add_heading('Graphs Replacement and Interventions', level=1)
        records, mastered, baseline_date = create_data_table(m, 'programs', medicaid_id)
        fill_table(document, 'programs', records, baseline_date, date)
        fill_with_graphics(document, 'programs', mastered, my_date, name_kid)

        document.add_page_break()

        address = os.path.join(os.getcwd(), os.path.join('DOCS', os.path.join('Monthlys', my_date)))
        if not os.path.isdir(address):
            os.mkdir(address)

        document.save(os.path.join(address, name_kid + '.docx'))


def create_data_table(m, table, medicaid_id):
    list_percent = []

    records = []
    list_i = m.list_behavior_or_program(table, str(medicaid_id))
    bl = m.list_baseline(table, str(medicaid_id), str(list_i[0]))
    baseline_date = str(bl[2]) + 'to' + str(bl[3])
    mastered = []

    for i in list_i:
        sto_string, status = list_sto(m, table, medicaid_id, i)

        if not status == 'MET':
            baseline_list = m.list_baseline(table, str(medicaid_id), str(i))
            aver = round(average_met(table, str(medicaid_id), str(i)), 2)
            if table is 'programs':
                baseline = str(baseline_list[0]) + '%'
                current_average = str(int(round(aver))) + '%'
                records.append((str(i), baseline, current_average, status, sto_string))
            else:
                baseline = str(baseline_list[0]) + ' incidents per ' + str(baseline_list[1]).lower()
                current_average = str(int(round(aver))) + ' incidents per week.'
                list_percent.append(((baseline_list[0] - aver) / baseline_list[0]) * 100)
                percent = str(int(round((((baseline_list[0] - aver) / baseline_list[0]) * 100), 2))) + '%'
                records.append((str(i), baseline, current_average, percent, status, sto_string))
        else:
            mastered.append(i)

    if table is 'behaviors':
        value = math.fsum(list_percent) / len(list_percent)
        records.append(('Total of Reduction (%)', '-', '-', '-', '-', str(value)))

    records = tuple(records)

    print 'mastered----'
    print mastered

    return records, mastered, baseline_date


def fill_table(document, name_table, records, baseline_date, date):

    if name_table is 'programs':
        count_cols = 5
        text = 'Replacements and Programs'
    else:
        count_cols = 6
        text = 'Maladaptive Behaviors'

    table = document.add_table(rows=1, cols=count_cols)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = text
    hdr_cells[1].text = 'Baseline Date:' + baseline_date
    hdr_cells[2].text = 'Current average per week (' + date + ')'
    if name_table is 'programs':
        hdr_cells[3].text = 'Status'
        i = 4
    else:
        hdr_cells[3].text = '% of Reduction(BL-average/ BL)*100'
        hdr_cells[4].text = 'Status'
        i = 5
    hdr_cells[i].text = 'Updated STO (' + date + ')'

    for value in records:
        row = table.add_row()
        row_cells = row.cells
        row_cells[0].text = value[0]  # name
        row_cells[1].text = value[1]  # bl
        row_cells[2].text = value[2]  # average
        row_cells[3].text = value[3]  # status if programs : percent
        row_cells[4].text = value[4]  # sto if programs : status
        if name_table is 'behaviors':
            row_cells[5].text = value[5]

    table.style = "TableGrid"


def fill_with_graphics(document, table_name, mastered, my_date, name_kid):
    name_kid = name_kid.split(' ')[0]

    dir_graphics = os.path.join(os.getcwd(), os.path.join('GRAPHICS', my_date))
    list_graphics = sorted_aphanumeric(os.listdir(dir_graphics))

    for graphic in list_graphics:
        var_bool = False
        if graphic.find(name_kid) is not -1 and graphic.find(table_name) is not -1:
            for met in mastered:
                if graphic.find(str(met)) is not -1:

                    print 'find graphic met----'
                    print str(met)

                    var_bool = True
            if not var_bool:
                document.add_picture(os.path.join(dir_graphics, graphic), width=Inches(6.5))


def create_excel(user, data_kid=None):
    m = Manager()
    my_date = datetime.datetime.now()
    my_date = my_date.replace(day=1) - timedelta(days=1)
    my_date = my_date.strftime("%B") + ' ' + str(my_date.year)
    address = os.path.join(os.getcwd(), os.path.join('DOCS', os.path.join('Excels', my_date)))

    if not os.path.isdir(address):
        os.mkdir(address)

    if data_kid is not None:
        kid_name = data_kid[0]
        medicaid_id = str(data_kid[1])
        list_kids = [medicaid_id + '-->' + kid_name]
    else:
        list_kids = m.list_kids(user)

    for kid in list_kids:
        medicaid_id = str(kid.split('-->')[0])
        kid_name = str(kid.split('-->')[1])

        name_file = kid_name + ' ' + medicaid_id + '.xlsx'
        if name_file in os.listdir(address):
            os.remove(os.path.join(address, name_file))
        # kid_name, last_name, kid_age, analyst
        data = m.kid_data(medicaid_id)[2]

        if not data:
            return False, 'Child not exist'

        workbook = xlsxwriter.Workbook(os.path.join(address, name_file))

        worksheet = workbook.add_worksheet('Name')
        worksheet.write(0, 0, data[0])
        worksheet.write(0, 1, data[1])
        worksheet.write(0, 2, data[2])
        worksheet.write(0, 3, medicaid_id)
        worksheet.write(0, 4, data[3])

        # region Sheet Behaviors
        worksheet = workbook.add_worksheet('Maladaptive Behaviors')
        list_beh = m.list_behavior_or_program('behaviors', medicaid_id)

        row = 0
        for beh in list_beh:
            create_excel_2(m, workbook, worksheet, str(beh), medicaid_id, row, 'Maladaptive Behaviors')
            row += 2
        # endregion

        # region Sheet Programs
        list_program = m.list_behavior_or_program('programs', medicaid_id)
        worksheet = workbook.add_worksheet('Programs')

        row = 0
        for program in list_program:
            create_excel_2(m, workbook, worksheet, str(program),  medicaid_id, row, 'Programs')
            row += 2
        # endregion

        worksheet = workbook.add_worksheet('lto_behaviors')
        expenses = ['Name', 'lto count', 'lto_time_count', 'lto_time']
        count_col = create_header_excel(workbook, worksheet, expenses)
        row = 1
        for beh in list_beh:
            b, s, data = m.lto_data('behaviors', medicaid_id, str(beh))
            if b:
                row = create_excel_3(data, worksheet, str(beh), count_col, row)
                # row += 1

        worksheet = workbook.add_worksheet('sto_behaviors')
        expenses = ['Name', 'sto count min', 'sto_count_max', 'sto_time_count', 'sto_time']
        count_col = create_header_excel(workbook, worksheet, expenses)
        row = 1
        for beh in list_beh:
            b, s, data = m.sto_data_without_string('behaviors', medicaid_id, str(beh))
            if b:
                row = create_excel_3(data, worksheet, str(beh), count_col, row)
                # row += 1

        worksheet = workbook.add_worksheet('lto_programs')
        expenses = ['Name', 'lto count', 'lto_time_count', 'lto_time']
        count_col = create_header_excel(workbook, worksheet, expenses)
        row = 1
        for program in list_program:
            b, s, data = m.lto_data('programs', medicaid_id, str(program))
            if b:
                row = create_excel_3(data, worksheet, str(program), count_col, row)
                # row += 1

        worksheet = workbook.add_worksheet('sto_programs')
        expenses = ['Name', 'sto count min', 'sto_count_max', 'sto_time_count', 'sto_time']
        count_col = create_header_excel(workbook, worksheet, expenses)
        row = 1
        for program in list_program:
            b, s, data = m.sto_data_without_string('programs', medicaid_id, str(program))
            if b:
                row = create_excel_3(data, worksheet, str(program), count_col, row)
                # row += 1

        workbook.close()

    return True, ''


def create_header_excel(workbook, worksheet, expenses):
    format1 = workbook.add_format({'font_color': 'red'})
    col = 0
    for v in expenses:
        worksheet.write(0, col, v, format1)
        col += 1
    return col


def create_excel_3(data, worksheet, beh_or_prog, count_col, row):
    # if is LTO
    if worksheet.name.startswith('lto'):
        col = 0
        worksheet.write(row, col, beh_or_prog)
        col += 1
        for v in data:
            worksheet.write(row, col, v)
            col += 1
            if col == count_col:
                break
        row += 1
    else:
        for l in data:
            col = 0
            worksheet.write(row, col, beh_or_prog)
            col += 1
            for v in l:
                worksheet.write(row, col, v)
                col += 1
                if col == count_col:
                    break
            row += 1
    return row


def create_excel_2(m, workbook, worksheet, beh_or_prog, medicaid_id, row, table_name):
    col = 0

    expenses = [[table_name, beh_or_prog]]
    if table_name == 'Maladaptive Behaviors':
        table_reference = 'behaviors'
    else:
        table_reference = 'programs'
    bl = m.list_baseline(table_reference, medicaid_id, beh_or_prog)
    # bl_count_time, bl_time, bl_begin_date, bl_end_date, exist_bl_end_date
    bl_count = -3
    if bl:
        if bl[4]:
            expenses.append([bl[2], bl[0]])
            expenses.append([bl[3], bl[0]])
            # expenses.append(['BL', 'BL'])
            # worksheet.write(row, col, bl[2])
            # col += 1
            # worksheet.write(row, col, bl[3])
            # col += 1
            # worksheet.write(row, col, 'BL')
            # col += 1
        else:
            expenses.append([bl[2], bl[0]])
            # expenses.append(['BL', 'BL'])
            # worksheet.write(row, col, bl[2])
            # col += 1
            # worksheet.write(row, col, 'BL')
            # col += 1

    # aqui escribo la fecha en rojo y su respectivo valor debajo
    expenses = tuple(expenses)
    format1 = workbook.add_format({'font_color': 'red'})
    for date, v in expenses:
        worksheet.write(row, col, date, format1)
        worksheet.write(row + 1, col, v)
        col += 1

    format1 = workbook.add_format({'font_color': 'green'})
    worksheet.write(row, col, 'BL', format1)
    worksheet.write(row + 1, col, 'BL', format1)
    col += 1
    expenses = []

    b, dict_dates_values = m.list_dates(table_reference, medicaid_id, beh_or_prog)
    list_dates = dict_dates_values.keys()

    list_tuple_value_comment = dict_dates_values.values()

    # list_dates = dt.datestr2num(list_dates)

    format1 = workbook.add_format({'num_format': 'mm/dd/yy', 'font_color': 'red'})
    for k, item in enumerate(list_dates):
        if list_tuple_value_comment[k][0] is -1:
            expenses.append([item, list_tuple_value_comment[k][1]])
        else:
            expenses.append([item, list_tuple_value_comment[k][0]])
    format2 = workbook.add_format({'font_color': 'red'})
    expenses = tuple(expenses)
    # Iterate over the data and write it out row by row.
    for date, value_comment in expenses:
        worksheet.write(row, col, date, format1)
        # worksheet.write(row, col, date)
        worksheet.write(row + 1, col, value_comment)
        col += 1


def create_monthly_2(username, medicaid_id=None, kid_name=None):
    m = Manager()

    ultimo_dia_mes_anterior = datetime.date.today().replace(day=1) + datetime.timedelta(days=-1)
    date = "%s/%s/%s" % (ultimo_dia_mes_anterior.month, ultimo_dia_mes_anterior.day, ultimo_dia_mes_anterior.year)
    if medicaid_id is not None:
        list_kids = [str(medicaid_id) + '-->' + str(kid_name)]
        list_medicaid = [medicaid_id]
    else:
        list_kids = m.list_kids(username)
        list_medicaid = m.list_medicaid_kid(username)

    for j, medicaid_id in enumerate(list_medicaid):
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        font.italic = True
        font.color.rgb = RGBColor(0, 0, 0)
        document.add_heading('BEHAVIOR ANALYST MONTHLY PROGRESS REPORT')

        table = document.add_table(rows=5, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        count_row = 0
        items = [
                    ['Assessment date', 'Initial assessment: 7/11/2017 Reassessment: 3/5/2018',
                     'DOB', '1/22/2004'],
                    ['Month/ Year', 'September/2018',
                     'Med#', '7944860172'],
                    ['Behavior Assistant', 'Angela M Fernandez',
                     'Mental Health Professional#', 'Davis, Safiyah'],
                    ['Authorization Initiation Date', '',
                     'Authorization Expiration Date', '']
        ]

        row_0 = table.rows[0].cells[0].merge(table.rows[0].cells[1].merge(table.rows[0].cells[2].merge(table.rows[0].cells[3])))
        text = 'Client:  John M Azoy'
        p = row_0.add_paragraph(text)
        p.alignment = WD_TABLE_ALIGNMENT.CENTER

        index_row = 1
        for item in items:
            row = table.rows[index_row]
            for index_col in range(0, 4):
                if index_col % 2 == 0:
                    p = row.cells[index_col].add_paragraph()
                    p.alignment = WD_TABLE_ALIGNMENT.CENTER
                    run = p.add_run(item[index_col])
                    run.bold = True
                else:
                    p = row.cells[index_col].add_paragraph()
                    p.alignment = WD_TABLE_ALIGNMENT.CENTER
                    run = p.add_run(item[index_col])
                    run.italic = False

            index_row += 1

        paragraph = document.add_paragraph()
        paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        p = document.add_paragraph()
        run = p.add_run('Summary: ')
        run.bold = True
        run.underline = True
        run = p.add_run(
            'During this month, we continued providing ABA services to John at his home and community, focusing '
            'on decrease the incidents of John`s maladaptive behaviors. Also, staff implemented different '
            'teaching'
            ' procedures and interventions to increase John`s social and basic living skills. Health and safety '
            'was monitored, there were no concerns. Caregiver does not report any changes in client`s routine '
            'during this month. Data collection of maladaptive behaviors and replacements was recorded by '
            'behavior technician and graphed/analyzed by behavior analyst. The behaviors will be maintained '
            'during the '
            'next evaluation period, with a schedule of short term objectives (STO) to decrease behavior`s '
            'incidents.'
            'John showed an increase in instructional goal programs. John showed a '
            'decrease'
            'in maladaptive behaviors incidents, mastering STO 2 for withdrawal of attention and '
            'argumentative/defiant behavior. Caregivers and Behavior assistant were trained in data collection '
            'and program implementation. Training activities specified on client`s behavior analysis support were'
            ' conducted. Regarding the parent training this month was mastered STO 4 for the three programs. '
            'We need to continue in observation and implementation of replacement programs. Analyst implemented '
            'program during the scheduled times, prescribed setting, correct materials gathered before beginning '
            'the program implementation, delivery of reinforces contingent upon correct behaviors and schedule '
            'reinforcement, and data was recorded appropriately. Analyst and behavior assistant will continue '
            'working on goals and objectives identified on behavior`s assessment.')
        run.bold = False
        run.italic = False

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        paragraph = document.add_paragraph()
        paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        paragraph = document.add_paragraph('Environmental medical/ biological, physical, social changes')
        paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY


        document.save(os.path.join(os.getcwd(), str(list_kids[j].split('-->')[1]) + '.docx'))


from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def table_print(block):
    table = block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text,'  ')
                #y.write(paragraph.text)
                #y.write('  ')
        print("\n")
        #y.write("\n")


def table_update(block, number_table):
    table = block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if number_table is 1:
                    table.cell(2, 1).text = ' '
                    p = table.rows[2].cells[1].add_paragraph()
                    p.alignment = WD_TABLE_ALIGNMENT.CENTER
                    p.add_run('December/2018')
                elif number_table is 3:
                    s = paragraph.text.split('(')[0]
                    if s == 'Current average per week ':
                        cell.text = ' '
                        p = cell.add_paragraph()
                        p.alignment = WD_TABLE_ALIGNMENT.CENTER
                        run = p.add_run('Current average per week (12/31/2018)')
                        run.bold = True
                elif number_table is 4 or number_table is 23:
                    s = paragraph.text.split('(')[0]
                    if s == 'Current average per week':
                        cell.text = ' '
                        p = cell.add_paragraph()
                        p.alignment = WD_TABLE_ALIGNMENT.CENTER
                        run = p.add_run('Current average per week')
                        run.bold = True
                        p.add_run('(12/31/2018)')
                    elif s == 'Updated STO':
                        cell.text = ' '
                        p = cell.add_paragraph()
                        p.alignment = WD_TABLE_ALIGNMENT.CENTER
                        run = p.add_run('Updated STO')
                        run.bold = True
                        p.add_run('(12/31/2018)')
                        p.line_spacing


if __name__ == '__main__':
    # create_excel('yessica')
    # create_monthly_2('yessica', '9460188605', 'Vanessa Victoria  Jr.')
    create_monthly('yessica')
    # document = Document('BA.docx')
    # count_table = 0
    # for block in iter_block_items(document):
    #     # print('found one')
    #     if not isinstance(block, Paragraph):
    #         count_table += 1
    #         print '<table>'
    #         table_print(block)
    #         if count_table is 1 or count_table is 23 or count_table is 3 or count_table is 4:
    #             table_update(block, count_table)
    #
    #
    #     # else:
    #     #     print block.text
    #     # print(block.text if isinstance(block, Paragraph) else '<table>')
    # print count_table
    # document.save('cambio a BA.docx')
    # table = document.tables[1]
    # table.cell(2, 1).text = ' '
    # p = table.rows[2].cells[1].add_paragraph()
    # p.alignment = WD_TABLE_ALIGNMENT.CENTER
    # run = p.add_run('December/2018')
    # run.bold = True
    # run.underline = True
    # document.save('Vanessa Victoria  Jr 2.docx')